"""DocumentParser — 文档解析器

支持格式：Markdown / YAML frontmatter / JSON / 纯文本 / PDF / ZIP / EPUB / HTML / DOCX
职责：读取文件 → 提取元数据 + 正文内容 → 返回结构化 Document 对象
编码检测：chardet（优先）→ utf-8 → gbk → latin-1 降级链
"""
from __future__ import annotations

import json
import logging
import os
import re
import xml.etree.ElementTree as ET
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

# chardet 可选依赖，不存在时降级到硬编码编码链
try:
    import chardet
    HAS_CHARDET = True
except ImportError:
    HAS_CHARDET = False

# PyMuPDF 可选依赖，不存在时跳过 PDF 解析
try:
    import fitz  # PyMuPDF
    HAS_FITZ = True
except ImportError:
    HAS_FITZ = False

# BeautifulSoup4 可选，用于 HTML 标签剥离
try:
    from bs4 import BeautifulSoup
    HAS_BS4 = True
except ImportError:
    HAS_BS4 = False

# python-docx 可选，用于 .docx 解析
try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


@dataclass
class Document:
    """解析后的文档结构"""
    file_path: str
    file_hash: str = ""
    file_size: int = 0
    create_time: str = ""
    modify_time: str = ""
    raw_content: str = ""
    frontmatter: dict = field(default_factory=dict)
    body: str = ""
    source_type: str = ""  # markdown / yaml / json / text / pdf / zip / epub / html / docx


class DocumentParser:
    """文档解析器：读取文件并提取结构化内容"""

    def parse(self, filepath: str) -> Optional[Document]:
        """解析单个文件，返回 Document 或 None（读取失败时）"""
        try:
            path = Path(filepath)
            if not path.exists():
                return None

            stat = path.stat()
            content = self._read_file(filepath)
            if content is None:
                return None

            doc = Document(
                file_path=filepath,
                file_size=stat.st_size,
                create_time=str(int(stat.st_ctime)),
                modify_time=str(int(stat.st_mtime)),
                raw_content=content,
            )

            # 提取 frontmatter（Markdown YAML 头）
            doc.frontmatter, doc.body = self._extract_frontmatter(content)

            # 判断来源类型
            doc.source_type = self._detect_type(filepath, content)

            return doc
        except Exception as e:
            logger.warning("Failed to parse %s: %s", filepath, e)
            return None

    def parse_batch(self, filepaths: list[str]) -> list[Document]:
        """批量解析文件"""
        results = []
        for fp in filepaths:
            doc = self.parse(fp)
            if doc:
                results.append(doc)
        return results

    def _read_file(self, filepath: str) -> Optional[str]:
        """读取文件内容，chardet 检测编码 → utf-8 → gbk → latin-1 降级链"""
        ext = Path(filepath).suffix.lower()

        # 专用格式路由（二进制格式走专用解析器）
        if ext == ".pdf":
            return self._read_pdf(filepath)
        if ext == ".zip":
            return self._read_zip(filepath)
        if ext == ".epub":
            return self._read_epub(filepath)
        if ext in (".html", ".htm"):
            return self._read_html(filepath)
        if ext == ".docx":
            return self._read_docx(filepath)

        # 通用文本格式：读取原始字节
        try:
            with open(filepath, "rb") as f:
                raw = f.read()
        except Exception as e:
            logger.warning("Cannot read file: %s: %s", filepath, e)
            return None

        if not raw:
            return ""

        # chardet 检测编码（优先）
        detected_enc = None
        if HAS_CHARDET and len(raw) > 100:
            result = chardet.detect(raw)
            if result and result.get("encoding") and result.get("confidence", 0) > 0.5:
                detected_enc = result["encoding"]

        # 编码降级链
        encodings = []
        if detected_enc:
            encodings.append(detected_enc)
        encodings.extend(["utf-8", "gbk", "latin-1"])

        for enc in encodings:
            try:
                return raw.decode(enc)
            except (UnicodeDecodeError, UnicodeError):
                continue

        logger.warning("Cannot decode file: %s", filepath)
        return None

    def _read_pdf(self, filepath: str) -> Optional[str]:
        """PDF 文件解析（需要 PyMuPDF）"""
        if not HAS_FITZ:
            logger.warning("PyMuPDF not installed, cannot parse PDF: %s", filepath)
            return None
        try:
            doc = fitz.open(filepath)
            text_parts = []
            for page in doc:
                text_parts.append(page.get_text())
            doc.close()
            return "\n".join(text_parts).strip()
        except Exception as e:
            logger.warning("PDF parse failed: %s: %s", filepath, e)
            return None

    def _read_zip(self, filepath: str) -> Optional[str]:
        """ZIP 文件解包：提取所有文本文件内容拼接"""
        try:
            with zipfile.ZipFile(filepath, "r") as zf:
                text_parts = []
                for name in sorted(zf.namelist()):
                    # 跳过目录和非文本文件
                    if name.endswith("/") or name.startswith("__MACOSX"):
                        continue
                    ext = Path(name).suffix.lower()
                    if ext not in (".txt", ".md", ".json", ".yaml", ".yml", ".py", ".js", ".csv"):
                        continue
                    try:
                        raw = zf.read(name)
                        # 简单编码检测
                        for enc in ("utf-8", "gbk", "latin-1"):
                            try:
                                text_parts.append(f"=== {name} ===\n{raw.decode(enc)}")
                                break
                            except UnicodeDecodeError:
                                continue
                    except Exception:
                        pass
                return "\n\n".join(text_parts) if text_parts else None
        except Exception as e:
            logger.warning("ZIP parse failed: %s: %s", filepath, e)
            return None

    def _read_epub(self, filepath: str) -> Optional[str]:
        """EPUB 电子书解析（零依赖：zipfile + xml.etree.ElementTree）

        流程：
        1. EPUB 本质是 ZIP，用 zipfile 打开
        2. 读 META-INF/container.xml 获取 .opf 文件路径
        3. 解析 OPF 的 <manifest> + <spine> 得到 XHTML 文件列表和阅读顺序
        4. 逐个解析 XHTML，用 ElementTree 提取 <body> 内纯文本
        """
        try:
            with zipfile.ZipFile(filepath, "r") as zf:
                # 1. 读 container.xml 获取 OPF 路径
                container_xml = zf.read("META-INF/container.xml")
                container_root = ET.fromstring(container_xml)
                # rootfile 的 full-path 属性是相对路径（相对于 EPUB 根目录）
                ns = {"ct": "urn:oasis:names:tc:opendocument:xmlns:container"}
                rootfile_el = container_root.find(".//ct:rootfile", ns)
                if rootfile_el is None:
                    rootfile_el = container_root.find(".//{urn:oasis:names:tc:opendocument:xmlns:container}rootfile")
                if rootfile_el is None:
                    logger.warning("EPUB: no rootfile found in container.xml: %s", filepath)
                    return None
                opf_path = rootfile_el.get("full-path", "")
                opf_dir = str(Path(opf_path).parent) if "/" in opf_path else ""

                # 2. 解析 OPF 文件
                opf_xml = zf.read(opf_path)
                opf_root = ET.fromstring(opf_xml)

                # 3. <manifest> 建立 id → href 映射
                manifest = {}
                for item in opf_root.iter("item"):
                    item_id = item.get("id", "")
                    href = item.get("href", "")
                    media = item.get("media-type", "")
                    if "html" in media or "xhtml" in media:
                        manifest[item_id] = href

                # 4. <spine> 获取阅读顺序
                spine_order = []
                for itemref in opf_root.iter("itemref"):
                    idref = itemref.get("idref", "")
                    if idref in manifest:
                        spine_order.append(manifest[idref])

                # 5. 逐个解析 XHTML 文件
                text_parts = []
                for href in spine_order:
                    # 拼接完整路径（相对于 OPF 目录）
                    full_href = f"{opf_dir}/{href}" if opf_dir else href
                    try:
                        xhtml = zf.read(full_href)
                        xhtml_str = xhtml.decode("utf-8", errors="replace")
                        # 用 ElementTree 解析
                        try:
                            xhtml_root = ET.fromstring(xhtml_str)
                        except ET.ParseError:
                            # XHTML 可能不是合法 XML，用正则兜底提取 <body> 内容
                            body_match = re.search(r'<body[^>]*>(.*?)</body>', xhtml_str, re.DOTALL | re.IGNORECASE)
                            if body_match:
                                raw_text = re.sub(r'<[^>]+>', '\n', body_match.group(1))
                                raw_text = re.sub(r'\n{3,}', '\n\n', raw_text).strip()
                                if raw_text:
                                    # 提取 <h1>-<h6> 作为章节标题
                                    title = href.split("/")[-1].replace(".xhtml", "").replace(".html", "")
                                    text_parts.append(f"=== {title} ===\n{raw_text}")
                            continue

                        # 从 <body> 提取文本
                        body_el = xhtml_root.find(".//{http://www.w3.org/1999/xhtml}body")
                        if body_el is None:
                            body_el = xhtml_root.find(".//body")
                        if body_el is not None:
                            raw_text = self._extract_element_text(body_el)
                            if raw_text.strip():
                                # 提取第一个标题作为章节名
                                title = ""
                                for tag in ("h1", "h2", "h3"):
                                    h_el = body_el.find(f".//{{http://www.w3.org/1999/xhtml}}{tag}")
                                    if h_el is None:
                                        h_el = body_el.find(f".//{tag}")
                                    if h_el is not None and h_el.text:
                                        title = h_el.text.strip()
                                        break
                                if not title:
                                    title = href.split("/")[-1].replace(".xhtml", "").replace(".html", "")
                                text_parts.append(f"=== {title} ===\n{raw_text.strip()}")
                    except Exception as e:
                        logger.debug("EPUB chapter read failed %s/%s: %s", opf_dir, href, e)

                return "\n\n".join(text_parts) if text_parts else None
        except Exception as e:
            logger.warning("EPUB parse failed: %s: %s", filepath, e)
            return None

    def _extract_element_text(self, element) -> str:
        """递归提取 XML/HTML 元素的纯文本"""
        parts = []
        if element.text:
            parts.append(element.text.strip())
        for child in element:
            parts.append(self._extract_element_text(child))
            if child.tail:
                parts.append(child.tail.strip())
        return " ".join(p for p in parts if p)

    def _read_html(self, filepath: str) -> Optional[str]:
        """HTML 文件解析：剥离标签提取纯文本"""
        try:
            with open(filepath, "rb") as f:
                raw = f.read()
        except Exception as e:
            logger.warning("Cannot read HTML: %s: %s", filepath, e)
            return None

        if not raw:
            return ""

        # 编码检测
        for enc in ("utf-8", "gbk", "latin-1"):
            try:
                html_str = raw.decode(enc)
                break
            except (UnicodeDecodeError, UnicodeError):
                continue
        else:
            logger.warning("Cannot decode HTML: %s", filepath)
            return None

        if HAS_BS4:
            soup = BeautifulSoup(html_str, "html.parser")
            # 移除 script 和 style 块
            for tag in soup(["script", "style", "nav", "footer", "header"]):
                tag.decompose()
            text = soup.get_text(separator="\n", strip=True)
        else:
            # 无 bs4 时：先移除 script/style 块，再剥标签
            html_str = re.sub(r'<(script|style|nav|footer|header)[^>]*>.*?</\1>', '', html_str, flags=re.DOTALL | re.IGNORECASE)
            text = re.sub(r'<[^>]+>', '\n', html_str)

        # 清理空行
        text = re.sub(r'\n{3,}', '\n\n', text).strip()
        return text if text else None

    def _read_docx(self, filepath: str) -> Optional[str]:
        """DOCX 文件解析（需要 python-docx）"""
        if not HAS_DOCX:
            logger.warning("python-docx not installed, cannot parse DOCX: %s", filepath)
            return None
        try:
            doc = docx.Document(filepath)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(paragraphs).strip() if paragraphs else None
        except Exception as e:
            logger.warning("DOCX parse failed: %s: %s", filepath, e)
            return None

    def _extract_frontmatter(self, content: str) -> tuple[dict, str]:
        """提取 YAML frontmatter 和正文

        格式：
        ---
        name: xxx
        description: xxx
        ---
        正文内容...
        """
        if not content.startswith("---"):
            return {}, content

        end = content.find("---", 3)
        if end == -1:
            return {}, content

        fm_text = content[3:end].strip()
        body = content[end + 3:].strip()

        # 简单 YAML 解析（不依赖 pyyaml）
        fm = {}
        for line in fm_text.split("\n"):
            line = line.strip()
            if not line or line.startswith("#"):
                continue
            if ":" in line:
                key, _, value = line.partition(":")
                key = key.strip()
                value = value.strip().strip('"').strip("'")
                if key:
                    fm[key] = value

        return fm, body

    def _detect_type(self, filepath: str, content: str) -> str:
        """根据扩展名和内容判断来源类型"""
        ext = Path(filepath).suffix.lower()
        type_map = {
            ".md": "markdown",
            ".markdown": "markdown",
            ".txt": "text",
            ".yaml": "yaml",
            ".yml": "yaml",
            ".json": "json",
            ".jsonl": "json",
            ".py": "code",
            ".js": "code",
            ".ts": "code",
            ".toml": "config",
            ".cfg": "config",
            ".ini": "config",
            ".pdf": "pdf",
            ".zip": "archive",
            ".epub": "epub",
            ".html": "html",
            ".htm": "html",
            ".docx": "docx",
        }
        if ext in type_map:
            return type_map[ext]
        # JSON 内容检测
        stripped = content.strip()
        if stripped.startswith(("{", "[")):
            try:
                json.loads(stripped)
                return "json"
            except (json.JSONDecodeError, ValueError):
                pass
        return "text"
